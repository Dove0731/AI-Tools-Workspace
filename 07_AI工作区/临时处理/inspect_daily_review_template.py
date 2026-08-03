from pathlib import Path
import json
import sys
from zipfile import ZipFile
from hashlib import sha256

from docx import Document


src = Path(sys.argv[1])
doc = Document(src)

items = []
body = doc.element.body
paragraph_index = 0
table_index = 0

for child in body.iterchildren():
    tag = child.tag.rsplit("}", 1)[-1]
    if tag == "p":
        p = doc.paragraphs[paragraph_index]
        items.append(
            {
                "type": "paragraph",
                "index": paragraph_index,
                "style": p.style.name if p.style else None,
                "text": p.text,
            }
        )
        paragraph_index += 1
    elif tag == "tbl":
        table = doc.tables[table_index]
        items.append(
            {
                "type": "table",
                "index": table_index,
                "rows": [
                    [cell.text.replace("\n", " / ") for cell in row.cells]
                    for row in table.rows
                ],
            }
        )
        table_index += 1

package_parts = []
with ZipFile(src) as zf:
    for info in sorted(zf.infolist(), key=lambda x: x.filename):
        data = zf.read(info.filename)
        package_parts.append(
            {
                "path": info.filename,
                "size": len(data),
                "sha256": sha256(data).hexdigest().upper(),
            }
        )

output = {
    "source": str(src),
    "source_sha256": sha256(src.read_bytes()).hexdigest().upper(),
    "paragraph_count": len(doc.paragraphs),
    "table_count": len(doc.tables),
    "items": items,
    "package_parts": package_parts,
}

print(json.dumps(output, ensure_ascii=False, indent=2))
