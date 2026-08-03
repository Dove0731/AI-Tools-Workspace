import json
import sys
import zipfile
import hashlib
import re
from collections import Counter
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


path = Path(sys.argv[1])
doc = Document(path)


def paragraph_info(p):
    fmt = p.paragraph_format
    run = next((r for r in p.runs if r.text.strip()), None)
    font = run.font if run else None
    return {
        "text": p.text,
        "style": p.style.name if p.style else None,
        "alignment": int(p.alignment) if p.alignment is not None else None,
        "space_before_pt": fmt.space_before.pt if fmt.space_before else None,
        "space_after_pt": fmt.space_after.pt if fmt.space_after else None,
        "line_spacing": fmt.line_spacing,
        "font_name": font.name if font else None,
        "font_size_pt": font.size.pt if font and font.size else None,
        "bold": font.bold if font else None,
    }


tables = []
for idx, table in enumerate(doc.tables, 1):
    grid = table._tbl.tblGrid
    grid_cols = []
    if grid is not None:
        for col in grid.gridCol_lst:
            grid_cols.append(col.get(qn("w:w")))
    rows = []
    for row in table.rows:
        rows.append([cell.text for cell in row.cells])
    tables.append({
        "index": idx,
        "rows": len(table.rows),
        "cols": len(table.columns),
        "style": table.style.name if table.style else None,
        "grid_twips": grid_cols,
        "sample_rows": rows[:3],
    })

sections = []
for idx, s in enumerate(doc.sections, 1):
    sections.append({
        "index": idx,
        "width_in": round(s.page_width.inches, 4),
        "height_in": round(s.page_height.inches, 4),
        "top_in": round(s.top_margin.inches, 4),
        "bottom_in": round(s.bottom_margin.inches, 4),
        "left_in": round(s.left_margin.inches, 4),
        "right_in": round(s.right_margin.inches, 4),
        "header_in": round(s.header_distance.inches, 4),
        "footer_in": round(s.footer_distance.inches, 4),
        "start_type": int(s.start_type),
    })

parts = []
with zipfile.ZipFile(path) as zf:
    document_xml = zf.read("word/document.xml").decode("utf-8")
    for info in sorted(zf.infolist(), key=lambda x: x.filename):
        data = zf.read(info.filename)
        parts.append({
            "path": info.filename,
            "size": len(data),
            "sha256": hashlib.sha256(data).hexdigest(),
        })

payload = {
    "sections": sections,
    "paragraphs": [paragraph_info(p) for p in doc.paragraphs],
    "tables": tables,
    "headers": [[p.text for p in s.header.paragraphs] for s in doc.sections],
    "footers": [[p.text for p in s.footer.paragraphs] for s in doc.sections],
    "parts": parts,
    "document_xml_tokens": {
        "fills": Counter(re.findall(r'w:fill="([^"]+)"', document_xml)),
        "colors": Counter(re.findall(r'w:color="([^"]+)"', document_xml)),
        "sizes_half_points": Counter(re.findall(r'w:sz w:val="([^"]+)"', document_xml)),
        "fonts": Counter(re.findall(r'w:(?:ascii|eastAsia|hAnsi)="([^"]+)"', document_xml)),
    },
}
print(json.dumps(payload, ensure_ascii=False, indent=2))
